"""
Docx Parser Module
Handles reading and extracting text from Word documents
"""

from typing import List, Dict, Tuple
from docx import Document
from docx.oxml.ns import qn
import os
import re


class DocxParser:
    """Parser for .docx files"""

    def __init__(self, file_path: str):
        """
        Initialize parser with docx file

        Args:
            file_path: Path to the .docx file
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        if not file_path.endswith('.docx'):
            raise ValueError("File must be a .docx file")

        self.file_path = file_path
        self.doc = Document(file_path)
        self.paragraphs_data = self._extract_paragraphs()

    def _get_all_paragraphs(self):
        """Helper to get all valid paragraphs sequentially (body + tables)"""
        paras = []
        for para in self.doc.paragraphs:
            paras.append(para)
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        paras.append(para)
        return paras

    def _get_para_text_with_formulas(self, element) -> Tuple[str, bool]:
        """
        Recursively extract text from paragraph elements, replacing complex MathXML with placeholders.
        Extracts plain text if it's a simple formula without complex nested structures.
        """
        parts = []
        has_formula = False
        formula_count = 0
        
        def walk(el):
            nonlocal has_formula, formula_count
            
            # Handle standard MathXML nodes
            if el.tag in (qn('m:oMathPara'), qn('m:oMath')):
                has_formula = True
                
                # Check if it contains simple m:t without complex structures
                m_texts = el.findall(f".//{qn('m:t')}")
                complex_tags = [qn('m:f'), qn('m:rad'), qn('m:sSub'), qn('m:sSup'), qn('m:sSubSup'), qn('m:m'), qn('m:d'), qn('m:eqArr')]
                has_complex = any(el.find(f".//{tag}") is not None for tag in complex_tags)
                
                if not has_complex and m_texts:
                    # Append raw latex or linear text directly
                    parts.append("".join([t.text for t in m_texts if t.text]))
                else:
                    formula_count += 1
                    parts.append(f"[[FORMULA_{formula_count}]]")
                return # skip walking deeper into formula
                
            if el.tag == qn('w:t'):
                parts.append(el.text or "")
            elif el.tag == qn('w:tab'):
                parts.append("\t")
            elif el.tag in (qn('w:br'), qn('w:cr')):
                parts.append("\n")
            else:
                for child in el:
                    walk(child)
                    
        walk(element)
        return "".join(parts), has_formula

    def _detect_drawing(self, para) -> bool:
        """Detect if a paragraph contains drawing, picture, OLE, or shape elements."""
        try:
            xml_str = para._element.xml
            return any(tag in xml_str for tag in [
                '<w:drawing', '<w:pict', '<w:object',
                '<v:shape', '<v:group', 'mc:AlternateContent'
            ])
        except Exception:
            return False

    @staticmethod
    def detect_layout(text: str) -> bool:
        """
        Detect if a paragraph is likely meant for layout/spacing only.
        
        Args:
            text: The paragraph text content
            
        Returns:
            True if it's primarily whitespace/formatting
        """
        if not text:
            return False
        stripped = text.strip()
        # Short equation labels like \t（3）, \t    （16）, (I) with leading whitespace
        if len(stripped) <= 8 and (text.startswith('\t') or text.startswith('   ')):
            return True
        if len(text) < 4:
            return False
        space_count = text.count(' ') + text.count('\t') * 4
        total_len = len(text)
        ratio = space_count / total_len if total_len > 0 else 0
        max_consec = max([len(m.group()) for m in re.finditer(r' {4,}', text)], default=0)
        return ratio > 0.5 or max_consec >= 4

    @staticmethod
    def detect_reference(text: str) -> bool:
        """
        Detect if a paragraph looks like a bibliographic reference entry.
        Returns True if it has a year and strong publication venue indicators.
        """
        if not text or len(text.strip()) < 15:
            return False
            
        t = text.strip()
        
        # MUST contain a 4-digit year (19xx or 20xx)
        has_year = bool(re.search(r'(?:19|20)\d{2}', t))
        if not has_year:
            return False
            
        # Chinese reference indicators
        cn_venue = bool(re.search(r'[《》]', t))
        cn_publisher = bool(re.search(r'出版社|工作论文|第\s*\d+\s*期|第\s*\d+\s*卷', t))
        cn_author_year = bool(re.search(r'[，,]\s*(?:19|20)\d{2}\s*[：:年]', t))
        
        # English reference indicators  
        en_journal = bool(re.search(r'Journal\s+of|Review\s+of|Quarterly|Economic[as]?|Financ|Working\s+Paper', t, re.IGNORECASE))
        en_venue = bool(re.search(r'pp?\.\s*\d+|Vol\.\s*\d+|No\.\s*\d+|\d+\s*[-\u2013]\s*\d+', t))
        en_author_year = bool(re.search(r'[,]\s*(?:19|20)\d{2}\s*[,]', t))
        en_publisher = bool(re.search(r'University\s+Press|Press[,.]|Publisher|Ed(?:s|itors)?\.', t, re.IGNORECASE))
        en_article = bool(re.search(r'["\u201c][^"\u201d]{3,100}["\u201d]', t))
        
        cn_score = sum([cn_venue, cn_publisher, cn_author_year])
        en_score = sum([en_journal, en_venue, en_author_year, en_publisher, en_article])
        
        return cn_score >= 1 or en_score >= 1

    def _extract_paragraphs(self) -> List[Dict]:
        """
        Extract all paragraphs with metadata, including tables.

        Returns:
            List of dicts with text, style, and formatting info
        """
        paragraphs = []
        all_paras = self._get_all_paragraphs()
        
        in_reference_section = False
        ref_heading_patterns = ['参考文献', 'References', 'Bibliography', 'REFERENCES']

        for idx, para in enumerate(all_paras):
            text, has_formula = self._get_para_text_with_formulas(para._element)
            contains_drawing = self._detect_drawing(para)
            
            is_ref = False
            style_name = para.style.name if para.style else "Normal"
            
            if text:
                stripped_text = text.strip()
                if not in_reference_section:
                    is_heading_style = 'Heading' in style_name or '标题' in style_name
                    is_ref_text = any(kw in stripped_text for kw in ref_heading_patterns)
                    is_short_ref_title = stripped_text in ref_heading_patterns or (is_ref_text and len(stripped_text) < 30)
                    
                    if (is_heading_style and is_ref_text) or (is_short_ref_title and not is_heading_style):
                        in_reference_section = True
                
                if in_reference_section:
                    is_ref = DocxParser.detect_reference(text)

            if not text:
                has_image_only = False
                has_text_content = False

                for run in para.runs:
                    if run._r is not None and ('<w:drawing>' in run._r.xml or '<w:pict>' in run._r.xml):
                        has_image_only = True
                    if run.text and run.text.strip():
                        has_text_content = True

                if has_image_only and not has_text_content:
                    paragraphs.append({
                        "index": idx,
                        "text": "",
                        "style": style_name,
                        "alignment": str(para.alignment) if para.alignment else None,
                        "runs": [],
                        "has_image": True,
                        "has_formula": False,
                        "contains_drawing": contains_drawing,
                        "is_layout": False,
                        "is_reference": False
                    })
                    continue
                else:
                    # Empty text but has drawing — still record it so it's protected
                    if contains_drawing:
                        paragraphs.append({
                            "index": idx,
                            "text": "",
                            "style": style_name,
                            "alignment": str(para.alignment) if para.alignment else None,
                            "runs": [],
                            "has_image": False,
                            "has_formula": False,
                            "contains_drawing": True,
                            "is_layout": False,
                            "is_reference": False
                        })
                    continue

            is_layout = DocxParser.detect_layout(text)

            para_info = {
                "index": idx,
                "text": text,
                "style": style_name,
                "alignment": str(para.alignment) if para.alignment else None,
                "runs": [],
                "has_image": False,
                "has_formula": has_formula,
                "contains_drawing": contains_drawing,
                "is_layout": is_layout,
                "is_reference": is_ref
            }

            for run in para.runs:
                if run.text:
                    para_info["runs"].append({
                        "text": run.text,
                        "bold": run.bold,
                        "italic": run.italic,
                        "underline": run.underline
                    })

            paragraphs.append(para_info)

        return paragraphs

    def get_all_text(self) -> str:
        """Get all text from document"""
        return "\n\n".join([p["text"] for p in self.paragraphs_data])

    def get_paragraphs(self, max_paragraphs: int = None) -> List[Dict]:
        """
        Get paragraphs, optionally limited to first N

        Args:
            max_paragraphs: Maximum number of paragraphs to return

        Returns:
            List of paragraph dictionaries
        """
        if max_paragraphs:
            return self.paragraphs_data[:max_paragraphs]
        return self.paragraphs_data

    def chunk_by_paragraph(self) -> List[Dict]:
        """
        Split text by paragraphs - recommended method for modern LLMs
        We no longer skip layout or image paragraphs here since the frontend
        explicitly allows the user to select them and `revision_writer` handles the fallback.

        Returns:
            List of chunks, each containing one paragraph with its index
        """
        chunks = []
        for para in self.paragraphs_data:
            # Yield every single parsed paragraph to ensure 1:1 mapping with the UI selection.
            # If a user selects an empty paragraph, we will process it or handle it cleanly.
            # CRITICAL: use para["index"] (real document position), NOT enumerate idx
            chunks.append({
                "text": para["text"],
                "paragraph_indices": [para["index"]],
                "paragraphs": [para],
                "is_reference": para.get("is_reference", False)
            })
        return chunks

    def chunk_full_document(self) -> List[Dict]:
        """
        Return entire document as a single chunk for consistency-focused proofreading
        Includes all paragraphs except purely empty text ones without visual content.

        Returns:
            List with single chunk containing all paragraphs
        """
        safe_paragraphs = [
            p for p in self.paragraphs_data
        ]
        
        has_reference_in_chunk = any(p.get("is_reference", False) for p in safe_paragraphs)

        # CRITICAL: use para["index"] (real document position) for all indices
        chunks = [{
            "text": "\n\n".join([p["text"] for p in safe_paragraphs]),
            "paragraph_indices": [p["index"] for p in safe_paragraphs],
            "paragraphs": safe_paragraphs,
            "is_reference": has_reference_in_chunk
        }]
        return chunks

    def get_chunk_for_testing(self, n_chunks: int = 2) -> List[Dict]:
        """
        Get first N text paragraphs for prompt testing
        Skips paragraphs containing images, drawings, or layout

        Args:
            n_chunks: Number of chunks to return

        Returns:
            List of chunks (text paragraphs only)
        """
        # chunk_by_paragraph already uses real doc indices
        all_chunks = self.chunk_by_paragraph()
        if not all_chunks:
            return []
        return all_chunks[:n_chunks]

    def get_document_stats(self) -> Dict:
        """
        Get document statistics

        Returns:
            Dictionary with document stats
        """
        total_chars = sum(len(p["text"]) for p in self.paragraphs_data)
        total_words = sum(len(p["text"].split()) for p in self.paragraphs_data)
        text_paragraphs = [p for p in self.paragraphs_data if not p.get("has_image", False)]
        image_paragraphs = [p for p in self.paragraphs_data if p.get("has_image", False) is True]

        return {
            "paragraphs": len(self.paragraphs_data),
            "text_paragraphs": len(text_paragraphs),
            "image_paragraphs": len(image_paragraphs),
            "characters": total_chars,
            "words": total_words,
            "estimated_tokens": total_chars // 3
        }

    @staticmethod
    def is_valid_docx(file_path: str) -> bool:
        """Check if file is a valid docx file"""
        if not os.path.exists(file_path):
            return False
        if not file_path.endswith('.docx'):
            return False

        try:
            # Try to open the file
            doc = Document(file_path)
            return True
        except Exception:
            return False
