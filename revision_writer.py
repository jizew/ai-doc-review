"""
Revision Writer Module
Handles writing revisions back to Word documents
"""

from typing import List, Dict, Optional
from docx import Document
from docx.shared import RGBColor
import difflib
import re
from docx.oxml.ns import qn
from docx_comment_utils import add_margin_comment
from docx_parser import DocxParser

try:
    from docx_revisions import RevisionDocument
    DOCX_REVISIONS_AVAILABLE = True
except ImportError:
    DOCX_REVISIONS_AVAILABLE = False

try:
    from docx_editor import DocxEditor
    DOCX_EDITOR_AVAILABLE = True
except ImportError:
    DOCX_EDITOR_AVAILABLE = False


class RevisionWriter:
    """Writer for Word document revisions"""

    def __init__(self, original_doc_path: str):
        """
        Initialize revision writer

        Args:
            original_doc_path: Path to original document
        """
        self.doc = Document(original_doc_path)
        self.original_doc_path = original_doc_path
        
        # Build a flat list of all paragraphs in the same order as in docx_parser
        self.all_paragraphs = []
        for para in self.doc.paragraphs:
            self.all_paragraphs.append(para)
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self.all_paragraphs.append(para)

    def _find_diff(self, original: str, revised: str) -> List[Dict]:
        """
        Find differences between original and revised text

        Args:
            original: Original text
            revised: Revised text

        Returns:
            List of differences with position and type
        """
        diffs = []
        matcher = difflib.SequenceMatcher(None, original, revised)

        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'replace':
                diffs.append({
                    "type": "replace",
                    "start": i1,
                    "end": i2,
                    "original": original[i1:i2],
                    "revised": revised[j1:j2]
                })
            elif tag == 'delete':
                diffs.append({
                    "type": "delete",
                    "start": i1,
                    "end": i2,
                    "original": original[i1:i2],
                    "revised": ""
                })
            elif tag == 'insert':
                diffs.append({
                    "type": "insert",
                    "start": i1,
                    "end": i1,
                    "original": "",
                    "revised": revised[j1:j2]
                })

        return diffs

    def write_comments(self, revisions: List[Dict], author: str = "AI Editor") -> Document:
        """
        Write revisions as Word comments

        Args:
            revisions: List of revision results from proofreading
            author: Comment author name

        Returns:
            Modified Document object
        """
        # Group revisions by paragraph index
        revisions_by_para = {}
        for rev in revisions:
            for idx in rev.get("paragraph_indices", []):
                if idx not in revisions_by_para:
                    revisions_by_para[idx] = []
                revisions_by_para[idx].append(rev)

        # Add comments to paragraphs
        for para_idx, para_revisions in revisions_by_para.items():
            if para_idx >= len(self.all_paragraphs):
                continue

            para = self.all_paragraphs[para_idx]

            # Check if there's an actual revision
            has_changes = any(
                rev.get("revised_text") != rev.get("original_text")
                for rev in para_revisions
            )

            if not has_changes:
                continue

            # Collect all comments
            comments_text = []
            for rev in para_revisions:
                if rev.get("comment"):
                    comments_text.append(rev["comment"])

                # Add suggested revision text if different from original
                if rev.get("revised_text") and rev.get("revised_text") != rev.get("original_text"):
                    # Find specific changes
                    diffs = self._find_diff(rev["original_text"], rev["revised_text"])
                    for diff in diffs:
                        if diff["type"] == "replace":
                            comments_text.append(
                                f"替换: '{diff['original']}' → '{diff['revised']}'"
                            )
                        elif diff["type"] == "delete":
                            comments_text.append(f"删除: '{diff['original']}'")
                        elif diff["type"] == "insert":
                            comments_text.append(f"插入: '{diff['revised']}'")

            # Add comment to paragraph
            if comments_text:
                comment_text = "\n".join([f"• {c}" for c in comments_text])
                try:
                    added = add_margin_comment(para, comment_text, author=author, initials=author[:2].upper())
                    if not added:
                        raise Exception("Native comment addition failed")
                except Exception as e:
                    print(f"Warning: Falling back to inline comments. Error: {e}")
                    # Fallback to inline if native fails
                    marker = para.add_run(f" [{author}]")
                    marker.font.color.rgb = RGBColor(0x42, 0xA5, 0xF5)
                    marker.font.italic = True

                    # Use the same paragraph object to avoid body/table index mismatch.
                    comment_para = para.insert_paragraph_before()
                    comment_run = comment_para.add_run(f"[{author} 批注]: {comment_text}")
                    comment_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
                    comment_run.font.italic = True

        return self.doc

    def add_header(self, model_name: str, author_name: str, target_doc: Optional[Document] = None) -> None:
        """
        Add or replace document header with AI Editor info
        
        Args:
            model_name: Name of the LLM used
            author_name: Name of the AI Editor
            target_doc: Optional Document to modify (defaults to self.doc)
        """
        from docx.shared import Pt
        doc_to_modify = target_doc if target_doc is not None else self.doc
        
        # Extract base model name (ignore provider prefix)
        short_model_name = model_name.split('/')[-1]
        header_text = f"{author_name} - {short_model_name}"
        
        # Access document sections
        for section in doc_to_modify.sections:
            # Update all header types (Primary, First Page, Even Page)
            for header_type in ['header', 'first_page_header', 'even_page_header']:
                try:
                    header = getattr(section, header_type)
                    
                    # Clear existing header content safely
                    for p in header.paragraphs:
                        # Instead of removing the paragraph object from its parent, 
                        # which can be unstable, we clear its runs.
                        for run in p.runs:
                            run.text = ""
                    
                    # If there are no paragraphs, add one. Otherwise use the first one.
                    if not header.paragraphs:
                        p = header.add_paragraph()
                    else:
                        p = header.paragraphs[0]
                    
                    # Set the text and style
                    p.text = header_text
                    p.alignment = 2  # Right alignment
                    
                    if p.runs:
                        run = p.runs[0]
                        run.font.size = Pt(9)
                        run.font.italic = True
                except (AttributeError, Exception):
                    continue

    def apply_fonts(self, cn_font: str = "宋体", en_font: str = "Times New Roman", target_doc: Optional[Document] = None) -> None:
        """
        Apply specified fonts to all runs in the document.
        Uses EastAsia for Chinese and ASCII/HAnsi for English.
        """
        doc_to_modify = target_doc if target_doc is not None else self.doc
        
        paragraphs_to_process = []
        for para in doc_to_modify.paragraphs:
            paragraphs_to_process.append(para)
        for table in doc_to_modify.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        paragraphs_to_process.append(para)

        # 1. Process all paragraphs (including those in tables)
        for para in paragraphs_to_process:
            for run in para.runs:
                self._set_run_font(run, cn_font, en_font)
        
        # 2. Process headers and footers
        for section in doc_to_modify.sections:
            for header_type in ['header', 'first_page_header', 'even_page_header']:
                try:
                    header = getattr(section, header_type)
                    for p in header.paragraphs:
                        for run in p.runs:
                            self._set_run_font(run, cn_font, en_font)
                except (AttributeError, Exception):
                    continue
            
            for footer_type in ['footer', 'first_page_footer', 'even_page_footer']:
                try:
                    footer = getattr(section, footer_type)
                    for p in footer.paragraphs:
                        for run in p.runs:
                            self._set_run_font(run, cn_font, en_font)
                except (AttributeError, Exception):
                    continue

    def _set_run_font(self, run, cn_font: str, en_font: str) -> None:
        """Helper to set Chinese and English fonts for a run"""
        run.font.name = en_font
        # Set the East Asia font via XML to ensure it applies to Chinese characters
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), cn_font)
        # Also ensure ascii/hAnsi are set to en_font
        rFonts.set(qn('w:ascii'), en_font)
        rFonts.set(qn('w:hAnsi'), en_font)

    def write_track_changes(self, revisions: List[Dict], author: str = "AI Editor") -> Optional[Document]:
        """
        Write revisions as Word track changes

        Args:
            revisions: List of revision results from proofreading
            author: Revision author name

        Returns:
            Modified Document object, or None if neither library is available
        """
        if DOCX_EDITOR_AVAILABLE:
            return self._write_track_changes_editor(revisions, author)
        elif DOCX_REVISIONS_AVAILABLE:
            return self._write_track_changes_revisions(revisions, author)
        else:
            print("Warning: Neither docx-editor nor docx-revisions installed. Track changes not supported.")
            return None

    def _write_track_changes_editor(self, revisions: List[Dict], author: str = "AI Editor") -> Optional[Document]:
        """Write track changes using docx-editor"""
        try:
            from docx_editor import DocxEditor

            editor = DocxEditor(self.original_doc_path)

            for rev in revisions:
                original_text = rev.get("original_text", "")
                revised_text = rev.get("revised_text", "")

                if not original_text or not revised_text:
                    continue

                if original_text == revised_text:
                    continue

                # Extract relevant paragraphs to check for drawings or formulas
                # Unlike revisions, editor operates string-globally, but we can peek the paras
                paras_to_check = []
                for idx in rev.get("paragraph_indices", []):
                    if idx < len(editor.doc.paragraphs):
                        paras_to_check.append(editor.doc.paragraphs[idx])
                
                # Formula safety fallback
                if "[[FORMULA_" in original_text:
                    print(f"Formula detected. Falling back to comment.")
                    for p in paras_to_check:
                        try:
                            add_margin_comment(p, f"建议修订 (含公式): {revised_text}", author=author, initials=author[:2].upper())
                        except Exception: pass
                    continue
                
                # Layout text safety fallback
                if DocxParser.detect_layout(original_text):
                    print("Layout structure detected. Falling back to comment.")
                    for p in paras_to_check:
                        try:
                            add_margin_comment(p, revised_text, author=author, initials=author[:2].upper())
                        except Exception: pass
                    continue
                
                # Drawing/Image safety fallback
                is_drawing = False
                for p in paras_to_check:
                    xml = p._element.xml
                    if "<w:drawing" in xml or "<v:shape" in xml or "<v:imagedata" in xml or "<w:pict" in xml:
                        is_drawing = True
                        break
                        
                if is_drawing:
                    print("Drawing structure detected. Falling back to comment.")
                    for p in paras_to_check:
                        try:
                            add_margin_comment(p, revised_text, author=author, initials=author[:2].upper())
                        except Exception: pass
                    continue

                diffs = self._find_diff(original_text, revised_text)

                for diff in reversed(diffs):
                    try:
                        if diff["type"] in ["replace", "delete"]:
                            editor.replace_tracked(
                                diff["original"],
                                diff["revised"],
                                author=author
                            )
                            editor.insert_tracked(
                                diff["revised"],
                                author=author
                            )
                    except Exception as e:
                        print(f"Warning: Failed to apply track change for diff (fallback triggered): {e}")
                        continue

            # Add explanatory margin comments for the track changes
            all_editor_paras = []
            for para in editor.doc.paragraphs:
                all_editor_paras.append(para)
            for table in editor.doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            all_editor_paras.append(para)

            for rev in revisions:
                if rev.get("comment"):
                    for para_idx in rev.get("paragraph_indices", []):
                        if para_idx < len(all_editor_paras):
                            try:
                                add_margin_comment(
                                    all_editor_paras[para_idx], 
                                    rev["comment"], 
                                    author=author, 
                                    initials=author[:2].upper()
                                )
                            except Exception as e:
                                print(f"Warning: Failed to add margin comment in track changes mode. Error: {e}")

            return editor.doc

        except Exception as e:
            import traceback
            traceback.print_exc()
            print(f"Error applying track changes with docx-editor: {e}")
            raise Exception(f"docx-editor error: {e}")

    def _write_track_changes_revisions(self, revisions: List[Dict], author: str = "AI Editor") -> Optional[Document]:
        """Write track changes using scoped paragraph replacement"""
        try:
            # We no longer strictly need RevisionDocument since we manipulate nodes directly,
            # but we keep rdoc structure to maintain compatibility and returning the Document.
            rdoc = RevisionDocument(self.original_doc_path)
            
            # handle API difference
            doc_obj = getattr(rdoc, 'document', getattr(rdoc, 'doc', None))

            # Add explanatory margin comments for the track changes
            all_rdoc_paras = []
            if doc_obj:
                for para in doc_obj.paragraphs:
                    all_rdoc_paras.append(para)
                for table in doc_obj.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                all_rdoc_paras.append(para)

            from docx_node_replace import paragraph_replace_tracked

            for rev in revisions:
                original_text = rev.get("original_text", "")
                revised_text = rev.get("revised_text", "")

                if not original_text or not revised_text:
                    continue

                if original_text == revised_text:
                    continue
                    
                para_indices = rev.get("paragraph_indices", [])
                
                # Formula safety fallback: replacing the entire paragraph will delete MathXML nodes.
                # If formulas are present, add a margin comment instead to preserve the original structure formatting.
                if "[[FORMULA_" in original_text:
                    print(f"Formula detected in paragraph {para_indices}. Falling back to comment to preserve formulas.")
                    for idx in para_indices:
                        if idx < len(all_rdoc_paras):
                            try:
                                add_margin_comment(
                                    all_rdoc_paras[idx], 
                                    f"建议修订 (含公式): {revised_text}", 
                                    author=author, 
                                    initials=author[:2].upper()
                                )
                            except Exception:
                                pass
                    continue
                    
                # Layout text safety fallback: replacing massive whitespace blocks destroys diagram alignments.
                if DocxParser.detect_layout(original_text):
                    print(f"Layout structure detected in paragraph {para_indices}. Falling back to comment.")
                    for idx in para_indices:
                        if idx < len(all_rdoc_paras):
                            try:
                                add_margin_comment(
                                    all_rdoc_paras[idx], 
                                    revised_text,
                                    author=author, 
                                    initials=author[:2].upper()
                                )
                            except Exception:
                                pass
                    continue

                # Drawing safety fallback: replacing the entire paragraph might delete <w:drawing> nodes.
                # If drawings are present, add a margin comment instead to preserve the original structure.
                is_drawing = False
                for idx in para_indices:
                    if idx < len(all_rdoc_paras):
                        para_xml = all_rdoc_paras[idx]._element.xml
                        if "<w:drawing" in para_xml or "<v:shape" in para_xml or "<v:imagedata" in para_xml or "<w:pict" in para_xml:
                            is_drawing = True
                            break
                            
                if is_drawing:
                    print(f"Drawing structure detected in paragraph {para_indices}. Falling back to comment to preserve layout.")
                    for idx in para_indices:
                        if idx < len(all_rdoc_paras):
                            try:
                                add_margin_comment(
                                    all_rdoc_paras[idx], 
                                    revised_text, # Just output the revised text as requested by user
                                    author=author, 
                                    initials=author[:2].upper()
                                )
                            except Exception:
                                pass
                    continue

                # Apply track changes ONLY to the specific paragraphs affected
                for idx in para_indices:
                    if idx < len(all_rdoc_paras):
                        try:
                            paragraph_replace_tracked(
                                all_rdoc_paras[idx],
                                original_text,
                                revised_text,
                                author=author
                            )
                        except Exception as e:
                            print(f"Warning: Failed to replace tracked paragraph {idx}. Error: {e}")
                            # Fallback to comment
                            try:
                                add_margin_comment(
                                    all_rdoc_paras[idx], 
                                    f"建议修订: {revised_text}", 
                                    author=author, 
                                    initials=author[:2].upper()
                                )
                            except Exception:
                                pass

            for rev in revisions:
                if rev.get("comment") and doc_obj:
                    for para_idx in rev.get("paragraph_indices", []):
                        if para_idx < len(all_rdoc_paras):
                            try:
                                add_margin_comment(
                                    all_rdoc_paras[para_idx], 
                                    rev["comment"], 
                                    author=author, 
                                    initials=author[:2].upper()
                                )
                            except Exception as e:
                                print(f"Warning: Failed to add margin comment in track changes mode. Error: {e}")

            return doc_obj

        except Exception as e:
            import traceback
            traceback.print_exc()
            print(f"Error applying track changes with docx-revisions: {e}")
            raise Exception(f"docx-revisions error: {e}")

    def save_document(self, output_path: str) -> None:
        """
        Save document to file

        Args:
            output_path: Output file path
        """
        self.doc.save(output_path)

    def save_track_changes_document(self, output_path: str, author: str = "AI Editor") -> bool:
        """
        Save document with track changes

        Args:
            output_path: Output file path
            author: Revision author name

        Returns:
            True if successful, False otherwise
        """
        if DOCX_EDITOR_AVAILABLE:
            try:
                from docx_editor import DocxEditor
                editor = DocxEditor(self.original_doc_path)
                editor.save(output_path)
                return True
            except Exception as e:
                print(f"Error saving track changes with docx-editor: {e}")
                return False

        elif DOCX_REVISIONS_AVAILABLE:
            try:
                rdoc = RevisionDocument(self.original_doc_path)
                rdoc.save(output_path)
                return True
            except Exception as e:
                print(f"Error saving track changes with docx-revisions: {e}")
                return False

        return False

    def create_summary(self, revisions: List[Dict]) -> str:
        """
        Create a summary of all revisions

        Args:
            revisions: List of revision results

        Returns:
            Summary text
        """
        total_edits = 0
        paragraphs_with_edits = 0
        comments = []

        for rev in revisions:
            if rev.get("revised_text") != rev.get("original_text"):
                total_edits += 1
                if rev.get("paragraph_indices"):
                    paragraphs_with_edits += len(set(rev["paragraph_indices"]))

            if rev.get("comment"):
                comments.append(rev["comment"])

        summary = f"""
AI 文档审校摘要
{"=" * 40}

处理段落总数: {len(revisions)}
修改段落数: {paragraphs_with_edits}
总修订次数: {total_edits}

详细批注:
"""
        for i, comment in enumerate(comments[:10], 1):  # Show first 10 comments
            summary += f"\n{i}. {comment}"

        if len(comments) > 10:
            summary += f"\n... 还有 {len(comments) - 10} 条批注"

        return summary
